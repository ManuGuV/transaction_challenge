import pandas
import docx
import os

#Function to create required folders to save files
def create_required_folders(current_directory):
    email_folder = current_directory + "/resources/email"
    db_folder = current_directory + "/resources/db"

    if not os.path.exists(email_folder):
        os.makedirs(email_folder)
    if not os.path.exists(db_folder):
        os.makedirs(db_folder)

#Function to generate word document to be sent
def generate_document(current_directory, total_balance, avg_debit, avg_credit, summary_per_month):
    # Create an instance of a word document
    doc = docx.Document()
    # Add headers to the document
    doc.add_heading('File Summary', 1)  
    doc.add_heading('General Summary', 3)
    for point in ['Total balance: ' + total_balance, 'Average debit amount: ' + avg_debit, 'Average credit amount: ' + avg_credit]:
        doc.add_paragraph(point, style='ListNumber')
    doc.add_heading('Monthly summary', 1)  
    # Create a table for each month that has a summary
    for summary in summary_per_month:
        doc.add_heading(summary[0], 2)
        table = doc.add_table(rows=3, cols=2)
        row = table.rows[0]
        row.cells[0].text = "Number of transactions"
        row.cells[1].text = str(summary[1][2])
        row = table.rows[1]
        row.cells[0].text = "Average debit amount"
        row.cells[1].text = str(summary[1][0])
        row = table.rows[2]
        row.cells[0].text = "Average credit amounts"
        row.cells[1].text = str(summary[1][1])
    doc.add_picture("logo.jpeg")
    # Now save the document to a location 
    doc.save(current_directory + "/resources/email/file_sumary.docx")

def send_email():
'''Empty function for the time being but this would probably have html code and the inclusion of the word
   document as attachment'''
    print("Email succesfully sent...")

#Function that gathers important information from csv
def process_file(file):
    months = ["January","February","March","April","May","June","July","August","September","October","November","December"]
    summary_per_month = {}
    transactions_df = pandas.read_csv(file, index_col=0)
    total_balance = sum(transactions_df['Transaction'])
    avg_debit = transactions_df.groupby('Id', as_index=False).apply(lambda x: x[x['Transaction'] < 0 ]).drop(columns=['Date'])
    avg_debit = avg_debit['Transaction'].mean(axis = 0)
    avg_credit = transactions_df.groupby('Id', as_index=False).apply(lambda x: x[x['Transaction'] > 0 ]).drop(columns=['Date'])
    avg_credit = avg_credit['Transaction'].mean(axis = 0)
    num_transactions = 0
    current_debit = 0
    current_credit = 0
    current_month = transactions_df.iloc[0]['Date'].split('/')[0]

    for index, row in transactions_df.iterrows():
        if current_month == row['Date'].split('/')[0]:
            if row['Transaction'] < 0:
                current_debit = current_debit + row['Transaction']
            else:
                current_credit = current_credit + row['Transaction']
            num_transactions +=1
        else:
            summary_per_month[months[int(current_month)-1]] = [current_debit, current_credit, num_transactions]
            num_transactions = 1
            if row['Transaction'] < 0:
                current_debit = row['Transaction']
                current_credit = 0
            else:
                current_credit = row['Transaction']
                current_debit = 0
            current_month = row['Date'].split('/')[0]
        summary_per_month[months[int(current_month)-1]] = [current_debit, current_credit, num_transactions]
    sorted_months = sorted(summary_per_month.items(),key =lambda x:months.index(x[0]))
    generate_document(current_directory, str(round(total_balance, 2)), str(round(avg_debit, 2)), str(round(avg_credit, 2)), sorted_months)
    print("Summary file has been generated...")


if __name__ == "__main__":
    current_directory = os.path.dirname(os.path.abspath(__file__))
    create_required_folders(current_directory)
    filename = current_directory + "/data/transactions.csv"
    process_file(filename)
    send_email()